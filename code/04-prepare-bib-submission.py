#!/usr/bin/env python3
"""
Convert PLOS Comp Bio v2 manuscript to a Briefings in Bioinformatics draft.

Changes:
1. Author byline: insert Alfredo Iacoangeli (2nd author), renumber affiliations
2. Drop Author Summary section (PLOS-specific; BiB has no equivalent)
3. Insert Key Points section after Abstract (BiB convention)
4. Apply numerical corrections from 05-compute-paper-stats.py verification:
   - §29 Mistral parse rate: 72/216 (33.3%) -> 12/216 (5.6%)
   - §29 other models parse range: "95 to 100%" -> "85 to 100%" (o3, o4-mini below 95%)
   - §32 A2 with-spec: 0.97 -> 0.96 (V3 method exact value)
   - §39 Gemini perfect-consistency: 33% -> 36% (V3 filter exact value)
   - §42 B1 no-spec: 0.70 -> 0.69; B3 0.62 -> 0.63; B3 with-spec 0.08 -> 0.09;
         B2 split: "1.00 in both" -> "0.99 no-spec, 1.00 with-spec"
   - §47 Replace DeepSeek V3 example (data shows opposite of claimed pattern)

References, formatting and word count are NOT touched in this pass; address those
once BiB confirms scope/handling.

Reads:  ../DOCS/PLOS-CompBiol-Manuscript-v2.docx
Writes: ../DOCS/BiB-Manuscript-v1.docx
"""
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent.parent
SRC = BASE / "DOCS" / "PLOS-CompBiol-Manuscript-v2.docx"
DST = BASE / "DOCS" / "BiB-Manuscript-v1.docx"

NEW_BYLINE = (
    "Manuel Corpas1,2,*, Alfredo Iacoangeli3,4,5,6, "
    "Segun Fatumo7,8, Heinner Guio7,9,10"
)

NEW_AFFILIATIONS = [
    "1. School of Life Sciences, University of Westminster, London, UK",
    "2. GENEQ Global, London, UK",
    "3. Department of Biostatistics and Health Informatics, King's College London, London, UK",
    "4. Department of Basic and Clinical Neuroscience, King's College London, London, UK",
    "5. Perron Institute for Neurological and Translational Science, Perth, Western Australia, Australia",
    "6. Biomedical Research Centre, South London and Maudsley NHS Foundation Trust, London, UK",
    "7. MRC/UVRI and LSHTM Uganda Research Unit, Entebbe, Uganda",
    "8. Precision Healthcare University Research Institute, Queen Mary University of London, London, UK",
    "9. Universidad de Ingenieria y Tecnologia (UTEC), Lima, Peru",
    "10. INBIOMEDIC Research and Technological Center, Lima, Peru",
]

KEY_POINTS = [
    "Frontier LLMs are unreliable for pharmacogenomics without specification: 9 frontier models averaged 92.4% phenotype accuracy, with worst-case 61% (Gemini 2.5 Flash) and 7 errors on the DPYD rs3918290 T/T fluorouracil-lethal case.",
    "A plain-text ClawBio SKILL.md specification raised mean phenotype accuracy to 100.0% (290 of 290 model-test-case-population combinations correct on all three runs), eliminated all DPYD lethal-case errors, and produced perfect 3-of-3 consistency across every evaluated combination.",
    "Specification-constrained execution disproportionately benefits non-European populations: 6 of 7 lethal-case errors occurred in Latin American (3) or East African (3) contexts under the no-specification condition; the specification closed this gap entirely.",
    "Tier B (contextual) scores fell with specification (B1 0.69 to 0.38; B3 0.63 to 0.09), a deliberate trade of confabulated narrative for machine-readable correctness, appropriate for clinical-grade deployment.",
    "Skill libraries operationalise the clinical-grade tier of the agentic genomics validation framework: auditable specification, version control, and decoupling of domain expertise from model capability.",
]


def find_para_by_text(doc, prefix):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            return i, p
    return None, None


def set_paragraph_text(p, new_text):
    """Replace text while preserving the first run's formatting; clear other runs."""
    if p.runs:
        p.runs[0].text = new_text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(new_text)


def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new <w:p> immediately after the given paragraph."""
    new_p_elem = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p_elem)
    np = Paragraph(new_p_elem, paragraph._parent)
    np.add_run(text)
    if style:
        np.style = paragraph.part.document.styles[style]
    return np


def remove_paragraph(p):
    p._element.getparent().remove(p._element)


def main():
    assert SRC.exists(), f"Source not found: {SRC}"
    doc = Document(str(SRC))

    # --- 1. Author byline ---
    paras = doc.paragraphs
    byline = paras[1]
    assert byline.text.strip().startswith("Manuel Corpas"), (
        f"Expected byline at paragraph 1, got: {byline.text[:80]}"
    )
    set_paragraph_text(byline, NEW_BYLINE)

    # --- 2. Affiliations ---
    # Original: paragraphs 2-7 hold affil 1-6 (Westminster, GENEQ, MRC/UVRI, QMUL, UTEC, INBIOMEDIC).
    # New layout requires inserting 4 KCL/Perron/BRC affiliations between paragraphs 3 and 4.
    affil_paras = [paras[i] for i in range(2, 8)]
    assert affil_paras[0].text.strip().startswith("1. School of Life Sciences"), affil_paras[0].text
    assert affil_paras[5].text.strip().startswith("6. INBIOMEDIC"), affil_paras[5].text

    set_paragraph_text(affil_paras[0], NEW_AFFILIATIONS[0])  # 1. Westminster
    set_paragraph_text(affil_paras[1], NEW_AFFILIATIONS[1])  # 2. GENEQ
    set_paragraph_text(affil_paras[2], NEW_AFFILIATIONS[6])  # 7. MRC/UVRI
    set_paragraph_text(affil_paras[3], NEW_AFFILIATIONS[7])  # 8. QMUL
    set_paragraph_text(affil_paras[4], NEW_AFFILIATIONS[8])  # 9. UTEC
    set_paragraph_text(affil_paras[5], NEW_AFFILIATIONS[9])  # 10. INBIOMEDIC

    # Insert 4 new affiliations after paragraph 3 (now "2. GENEQ").
    # addnext places the new element directly after; insert in reverse so order is preserved.
    insert_paragraph_after(affil_paras[1], NEW_AFFILIATIONS[5])  # 6. BRC Maudsley
    insert_paragraph_after(affil_paras[1], NEW_AFFILIATIONS[4])  # 5. Perron
    insert_paragraph_after(affil_paras[1], NEW_AFFILIATIONS[3])  # 4. KCL Neuro
    insert_paragraph_after(affil_paras[1], NEW_AFFILIATIONS[2])  # 3. KCL Biostat

    # --- 3. Drop Author Summary heading + content ---
    idx, heading_p = find_para_by_text(doc, "Author Summary")
    assert heading_p is not None, "Could not find Author Summary heading"
    content_p = doc.paragraphs[idx + 1]
    assert content_p.style.name != "Heading 1", (
        f"Expected non-heading content after Author Summary, got: {content_p.style.name}"
    )
    remove_paragraph(content_p)
    remove_paragraph(heading_p)

    # --- 4. Key Points section after Abstract content ---
    abs_idx, abs_heading = find_para_by_text(doc, "Abstract")
    assert abs_heading is not None, "Could not find Abstract heading"
    abs_content = doc.paragraphs[abs_idx + 1]
    assert abs_content.style.name != "Heading 1", "Abstract content paragraph not found as expected"

    # Bullets: insert in reverse so final order matches KEY_POINTS list.
    for bullet in reversed(KEY_POINTS):
        insert_paragraph_after(abs_content, "• " + bullet)

    # Heading immediately after Abstract content.
    insert_paragraph_after(abs_content, "Key Points", style="Heading 1")

    # --- 5. Apply numerical corrections from stats verification (05-compute-paper-stats.py) ---
    apply_corrections(doc)

    # --- save ---
    doc.save(str(DST))
    print(f"Wrote: {DST}")


# Each correction: (find_substring, full_replacement_text). The find substring
# must uniquely identify the paragraph and remain present in the source.
CORRECTIONS = [
    # §29 Overview: Mistral parse rate + other-models range
    (
        "Mistral Large 2 produced parseable output for only",
        "We evaluated 1,944 model-test-case-population-condition-run combinations "
        "(9 models, 12 test cases, 3 population contexts, 2 conditions, 3 runs). "
        "Of these, 1,681 (86.5%) produced parseable output. One API error occurred. "
        "Mistral Large 2 produced parseable output for only 12 of 216 runs (5.6%), "
        "reflecting persistent format non-compliance rather than clinical errors; "
        "all other models parsed at 85 to 100%. Results below exclude unparsed "
        "responses unless otherwise noted.",
    ),
    # §32 Tier A overview: A2 0.87 -> 0.96
    (
        "rose from a cross-model mean of 0.87 to 0.97",
        "Table 1 shows mean Tier A scores by model and condition across all three "
        "population contexts. The specification improved all three clinical dimensions, "
        "with the largest effect on drug recommendation (A2), which rose from a "
        "cross-model mean of 0.87 to 0.96.",
    ),
    # §42 Tier B: rounding fixes (B1 0.70->0.69, B3 0.62->0.63, B3 0.08->0.09, B2 1.00->0.99/1.00)
    (
        "Dataset Specificity (B1) dropped from a mean of 0.70 to 0.38",
        "Table 2 shows Tier B scores. An unexpected finding was that Tier B scores "
        "decreased with specification for two of three dimensions: Dataset Specificity "
        "(B1) dropped from a mean of 0.69 to 0.38, and Domain Grounding (B3) from 0.63 "
        "to 0.09. This reflects the specification constraining output to the exact "
        "5-line format, which reduces the verbosity and domain elaboration that models "
        "produce when unconstrained. Reasoning Chain completeness (B2) was near 1.00 "
        "in both conditions (0.99 without specification, 1.00 with specification).",
    ),
    # §39 Consistency analysis: Gemini perfect-consistency rate 33% -> 36%
    (
        "Gemini 2.5 Flash achieved perfect consistency on only 33%",
        "Figure 1 shows the consistency heatmap across all model-test-case "
        "combinations (aggregated across populations). Without specification, 241 of "
        "275 model-test-case-population combinations (87.6%) achieved perfect "
        "consistency (3 of 3 correct runs across all three runs). Model-level "
        "consistency without specification varied substantially (Figure 2, right "
        "panel). Gemini 2.5 Flash achieved perfect consistency on only 36% of "
        "combinations. With specification, all 290 evaluable combinations achieved "
        "perfect consistency (100%), with zero stochastic failures.",
    ),
    # §47 Population effect: drop wrong DeepSeek V3 example, reframe around DPYD
    (
        "DeepSeek V3 dropped from 100% to approximately 60%",
        "Figure 2 shows phenotype accuracy (A1) by population context. Without "
        "specification, models showed a downward trend from European (93% mean A1) "
        "through Latin American (92%) to East African (92%) contexts. This aggregate "
        "difference is small and is driven primarily by Gemini 2.5 Flash, which "
        "dropped from 67% (European) to 53% (East African); other frontier models "
        "showed minimal aggregate population dependence. The population-dependent "
        "reliability emerged most clearly on safety-critical cases, as the DPYD "
        "lethal-case results above demonstrate.",
    ),
]


def apply_corrections(doc):
    applied = []
    for find_str, replacement in CORRECTIONS:
        target = None
        for p in doc.paragraphs:
            if find_str in p.text:
                target = p
                break
        assert target is not None, f"Could not find paragraph containing: {find_str!r}"
        set_paragraph_text(target, replacement)
        applied.append(find_str[:50])
    print(f"Applied {len(applied)} numerical corrections:")
    for a in applied:
        print(f"  - {a}...")


if __name__ == "__main__":
    main()
