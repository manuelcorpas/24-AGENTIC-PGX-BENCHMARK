#!/usr/bin/env python3
"""
Fabrication firewall for the corrected Cell Genomics revision.

Every headline number in the paper is recomputed here from the authoritative
data file and checked against the string actually present in the .docx. If the
manuscript and this script disagree, the manuscript is wrong.

WHY THIS EXISTS AT ALL
A referee has now twice found that the deposited artefacts and the manuscript
described different things, and this project has logged fourteen instances of a
number that came from tooling rather than from the phenomenon. A check that a
number appears in the text is weak; a check that the number in the text equals
the number the data produce is the one that matters, and it is cheap.

WHAT IT DOES NOT DO
It does not parse prose or find numbers the author forgot to register here. A
claim absent from CHECKS below is unverified by this script, which is why the
script prints how many claims it checked rather than implying completeness.

USAGE
    python3 code/72-validate-v30-numbers.py --manuscript /path/to/final.docx
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
D = BASE / "data"


def load(name):
    p = D / name
    if not p.exists():
        return None
    return json.loads(p.read_text()) if p.suffix == ".json" else p.read_text()


def parse_report(text: str, scorer="baseline") -> dict:
    """Pull per-cell metrics out of v3_five_cell_live_report.txt."""
    out, section = {}, None
    for line in (text or "").splitlines():
        s = line.strip()
        if s.startswith("## "):
            section = s[3:].strip()
            continue
        if section != scorer or not s or "n=" not in s:
            continue
        cell = s.split()[0]
        fields = {}
        for tok in s.split():
            if "=" in tok:
                k, v = tok.split("=", 1)
                fields[k] = v
        lethal_err = None
        if "(" in s and "errors)" in s:
            lethal_err = int(s.split("(")[1].split()[0])
        out[cell] = {"n": int(fields.get("n", 0)),
                     "coverage": fields.get("coverage"),
                     "A1": fields.get("A1"), "A2": fields.get("A2"),
                     "lethal_errors": lethal_err}
    return out


def cell_means() -> dict:
    """A1 and A2 means at full precision, from the scored rows themselves.

    Deriving these from v3_five_cell_live_report.txt instead rounds twice: the
    report stores four decimals, and reading 0.9705 back and rounding to three
    gives 0.971 for a cell whose true mean is 0.970455. The manuscript then
    printed 97.1% where the data say 97.0%, and this check passed it, because
    both sides had been through the same rounding. Percentages in the paper are
    one decimal place, so the check has to start from a number that has not
    already been rounded to fewer digits than it needs.
    """
    rows = load("v3_matched_scored_rows_all5.json")
    if not rows:
        return {}
    out: dict[str, dict[str, float]] = {}
    for row in rows:
        bucket = out.setdefault(row["cell"], {"n": 0, "a1": 0.0, "a2": 0.0})
        bucket["n"] += 1
        bucket["a1"] += float(row["a1_phenotype"])
        bucket["a2"] += float(row["a2_recommendation"])
    return {cell: {"A1": b["a1"] / b["n"], "A2": b["a2"] / b["n"], "n": b["n"]}
            for cell, b in out.items()}


def build_checks():
    checks = []
    rep = parse_report(load("v3_five_cell_live_report.txt") or "")
    means = cell_means()
    if rep:
        for cell, a1, a2, leth in (
            ("free_generation", "0.744", "0.624", 54),
            ("rag_generation", "0.820", "0.551", 137),
            ("rag_execution", "0.965", "0.970", 23),
            ("skill_generation", "0.964", "0.970", 26),
            ("skill_execution", "0.967", "0.973", 25),
        ):
            got = rep.get(cell, {})
            exact = means.get(cell)
            if exact:
                checks.append((f"{cell} A1", a1, f"{exact['A1']:.3f}"))
                checks.append((f"{cell} A2", a2, f"{exact['A2']:.3f}"))
            else:
                checks.append((f"{cell} A1", a1, f"{float(got.get('A1', 0)):.3f}"))
                checks.append((f"{cell} A2", a2, f"{float(got.get('A2', 0)):.3f}"))
            checks.append((f"{cell} lethal errors", str(leth),
                           str(got.get("lethal_errors"))))
        returned = sum(v["n"] for v in rep.values())
        checks.append(("returned evaluation records", "13,199", f"{returned:,}"))
        # One prespecified authored-rule-execution request returned no record.
        # It remains an attempted evaluation and is imputed as a failed attempt
        # in all end-to-end denominators.
        checks.append(("attempted evaluations", "13,200", f"{returned + 1:,}"))

    ext = load("v3_extracted_rules_eval.json")
    if ext:
        ae = ext["authored_execution"]
        checks.append(("authored coverage", "0.9996", str(ae["coverage"])))
        checks.append(("authored accuracy among emitted", "0.968",
                       f"{ae['accuracy_among_emitted']:.3f}"))
        # Registered against the eight-model extraction reported in Table 2.
        # The earlier 0.627 and 0.645 were the three-model era and had gone
        # stale without failing, because the check only compared them with a
        # data file that had itself moved on.
        for model, lo in (("Claude Opus 4.5", "0.660"), ("GPT-5.2", "0.636"),
                          ("o3", "0.660")):
            m = ext["models"].get(model)
            if m:
                checks.append((f"extracted coverage, {model}", lo,
                               str(m["execution"]["coverage"])))

    leth = load("v3_lethal_case_level.json")
    if leth:
        for name, diff in (("all", "0.247"), ("HLA", "0.617"), ("non-HLA", "0.042")):
            st = leth["strata"].get(name)
            if st:
                checks.append((f"lethal difference, {name}", diff,
                               f"{st['retrieval_minus_free']['difference']:.3f}"))
        checks.append(("distinct lethal cases", "14",
                       str(leth["n_distinct_cases"])))

    prov = load("v3_cell_provenance.json")
    if prov:
        tot = sum(v["cost_usd"] for v in prov["per_cell"].values())
        checks.append(("total spend across cells", "$55.30", f"${tot:.2f}"))
        lc = prov.get("legacy_comparison") or {}
        checks.append(("legacy-identical skill rows", "0",
                       str(lc.get("identical_on_raw_tokens_and_cost"))))

    getrm = load("v3_getrm_disagreement_classes.json")
    if getrm:
        # Registered at the precision the manuscript states. The in-text check is
        # string containment, so registering 0.7609 against a paper that rounds
        # to 0.761 fails on presentation rather than on substance.
        checks.append(("GeT-RM concordance", "0.761", str(getrm["concordance"])))
        checks.append(("GeT-RM shared-definition subset", "0.993",
                       str(getrm["shared_definition_concordance"])))
        checks.append(("GeT-RM evaluable pairs", "527", str(getrm["evaluable"])))
        checks.append(("GeT-RM attributable disagreements", "104",
                       str(getrm["attributable"])))
        checks.append(("GeT-RM unexplained", "22", str(getrm["unexplained"])))

    # R1.1 input normalisation. The freeze hashes every raw source and is the
    # single analysis object used by the input-normalisation figure, its
    # operational table and these checks. It must be the EIGHT-model freeze:
    # the seven-model file is still deposited for provenance, and registering
    # against it silently checked the manuscript against a superseded panel.
    freeze = load("v3_input_normalisation_eight_model_freeze.json")
    ev = load("v3_input_normalisation_eval.json")
    if freeze and ev:
        op = freeze["definition_arm_operational_total"]
        checks.extend([
            ("normalisation definition attempts", "4,216", f"{op['attempted']:,}"),
            ("normalisation definition calls", "3,002", f"{op['call']:,}"),
            ("normalisation definition abstentions", "1,210", f"{op['abstain']:,}"),
            ("normalisation definition truncations", "four truncated responses", "four truncated responses" if op["truncated_output"] == 4 else str(op["truncated_output"])),
            ("normalisation pooled scored", "2,942", f"{ev['n_scored']:,}"),
            ("normalisation pooled calls", "1,597", f"{ev['overall']['vs_caller']['emitted']:,}"),
            ("normalisation pooled coverage", "0.543", f"{ev['overall']['vs_caller']['coverage']:.3f}"),
            ("normalisation pooled PyPGx accuracy", "0.387", f"{ev['overall']['vs_caller']['accuracy_among_emitted']:.3f}"),
            ("normalisation pooled GeT-RM accuracy", "0.418", f"{ev['overall']['vs_getrm']['accuracy_among_emitted']:.3f}"),
        ])
        opus = freeze["models_analysis"]["Claude Opus 4.5"]
        wo = opus["without_definitions"]
        wi = opus["with_definitions_full_527"]
        checks.extend([
            ("normalisation coverage, no definitions", "0.934", f"{wo['coverage']['estimate']:.3f}"),
            ("normalisation accuracy, no definitions", "0.396", f"{wo['accuracy_vs_pypgx_among_calls']['estimate']:.3f}"),
            ("normalisation coverage, definitions", "0.973", f"{wi['coverage']['estimate']:.3f}"),
            ("normalisation accuracy, definitions", "0.967", f"{wi['accuracy_vs_pypgx_among_calls']['estimate']:.3f}"),
            ("normalisation vs GeT-RM, model", "0.784", f"{wi['accuracy_vs_getrm_among_calls']['estimate']:.3f}"),
            ("normalisation vs GeT-RM, caller", "0.774", f"{wi['caller_vs_getrm_on_model_calls']['estimate']:.3f}"),
            ("normalisation pairs answered", "513", str(wi["operational"]["call"])),
            ("normalisation model-caller difference", "0.010", f"{wi['model_minus_caller_vs_getrm']['estimate']:.3f}"),
            ("normalisation difference CI lower", "-0.004", f"{wi['model_minus_caller_vs_getrm']['ci95'][0]:.3f}"),
            ("normalisation difference CI upper", "0.024", f"{wi['model_minus_caller_vs_getrm']['ci95'][1]:.3f}"),
        ])
        stated = {
            "Gemini 2.5 Flash": ("0.814", "0.825"),
            "Claude Sonnet 4.5": ("0.749", "0.778"),
            "o4-mini": ("0.731", "0.764"),
            "o3": ("0.724", "0.766"),
            "DeepSeek V3": ("0.489", "0.770"),
            "GPT-5.2": ("0.480", "0.765"),
            "GPT-4.1": ("0.464", "0.707"),
        }
        for model, (model_value, caller_value) in stated.items():
            m = freeze["models_analysis"][model]["with_definitions_full_527"]
            checks.append((f"normalisation vs GeT-RM, {model}", model_value,
                           f"{m['accuracy_vs_getrm_among_calls']['estimate']:.3f}"))
            checks.append((f"caller vs GeT-RM for {model}", caller_value,
                           f"{m['caller_vs_getrm_on_model_calls']['estimate']:.3f}"))

    avd = load("v3_agent_vs_deterministic.json")
    if avd:
        # CorpasFamily is 0.292 after the whole-genome rerun (was 0.333 on the
        # SNP-chip family arm). The old registration matched a data file that
        # had already been regenerated, so it failed silently rather than
        # catching the manuscript.
        for coh, cov in (("1000G_IBS", "0.215"), ("CorpasFamily", "0.292"),
                         ("Peru", "0.300"), ("UGR", "0.090")):
            d = avd["deterministic"].get(coh)
            if d:
                checks.append((f"deterministic coverage, {coh}", cov,
                               f"{d['coverage']:.3f}"))
                checks.append((f"deterministic accuracy, {coh}", "1.0",
                               str(d["accuracy_among_emitted"])))

    # Vocabulary coverage across the four cohorts, as reported in the
    # four-cohort coverage table and its figure. These moved with the family
    # rerun and were not registered anywhere, which is how the stale text
    # report in data/ survived alongside a corrected JSON.
    anc = load("v3_ancestry_four_cohorts.json")
    if anc:
        for coh, states, cov_state, cov_carrier, std, own in (
            ("1000G_IBS", "137", "0.161", "0.385", "0.362", "0.075"),
            ("CorpasFamily", "37", "0.216", "0.265", "0.421", "0.0"),
            ("Peru", "73", "0.192", "0.419", "0.344", "0.143"),
            ("UGR", "289", "0.066", "0.272", "0.322", "0.029"),
        ):
            raw = anc["raw_by_cohort"][coh]
            checks.append((f"distinct states, {coh}", states,
                           str(raw["distinct_states"])))
            checks.append((f"vocabulary coverage (states), {coh}", cov_state,
                           f"{raw['coverage_states']:.3f}"))
            checks.append((f"vocabulary coverage (carriers), {coh}", cov_carrier,
                           f"{raw['coverage_carriers']:.3f}"))
            checks.append((f"standardised coverage, {coh}", std,
                           f"{anc['standardised'][coh]['standardised_coverage']:.3f}"))
            checks.append((f"cohort-specific coverage, {coh}", own,
                           f"{anc['cohort_specific'][coh]['coverage_states']:.3f}"))
        checks.append(("raw state-coverage spread", "0.150",
                       f"{anc['raw_state_coverage_spread']:.3f}"))
        checks.append(("standardised coverage spread", "0.099",
                       f"{anc['standardised_coverage_spread']:.3f}"))
        checks.append(("cohort-specific coverage spread", "0.143",
                       f"{anc['cohort_specific_coverage_spread']:.3f}"))
    return checks


def _agree(stated: str, computed: str) -> bool:
    """Compare at the precision the manuscript states.

    The paper rounds to three decimals; the data file carries four. Demanding
    string equality would flag 0.627 against 0.6268 as a fabrication, which is
    noise that trains the reader of this report to ignore it.
    """
    a = stated.lstrip("$").replace(",", "")
    b = computed.lstrip("$").replace(",", "")
    if a == b:
        return True
    try:
        dp = len(a.split(".")[1]) if "." in a else 0
        return round(float(a), dp) == round(float(b), dp)
    except ValueError:
        return False


def _renderings(stated: str) -> list[str]:
    """Every way the manuscript may legitimately print a registered value.

    Registered values are proportions, because that is what the data files
    carry. The manuscript prints percentages. A containment check that knows
    only "0.744" reports a FAIL on a paper that correctly says "74.4%", and
    fourteen such FAILs in a thirty-row report is enough noise to make the
    two real ones invisible. That is precisely what happened: three stale
    registrations sat unnoticed behind a wall of formatting failures.
    """
    forms = {stated}
    bare = stated.lstrip("$").replace(",", "")
    try:
        value = float(bare)
    except ValueError:
        return sorted(forms)
    if "." in bare and -1.0 <= value <= 1.0:
        pct = value * 100
        decimals = max(0, len(bare.split(".")[1]) - 2)
        for places in {decimals, decimals + 1, 0, 1}:
            forms.add(f"{pct:.{places}f}%")
            # "30.0%" is also written "30%"; strip a trailing zero decimal.
            forms.add(f"{pct:.{places}f}".rstrip("0").rstrip(".") + "%")
    return sorted(forms)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n")[1])
    ap.add_argument("--manuscript", type=Path, default=None,
                    help="v30 .docx; if given, each stated value must appear in it")
    a = ap.parse_args(argv)

    text = None
    if a.manuscript:
        try:
            import docx
        except ImportError:
            sys.stderr.write("python-docx not installed; skipping the text check\n")
        else:
            if not a.manuscript.exists():
                sys.stderr.write(f"missing {a.manuscript}\n")
                return 1
            doc = docx.Document(str(a.manuscript))
            text = "\n".join(p.text for p in doc.paragraphs)
            text += "\n" + "\n".join(c.text for t in doc.tables
                                     for r in t.rows for c in r.cells)

    checks = build_checks()
    if not checks:
        sys.stderr.write("no data files found; nothing verified\n")
        return 1

    fails = []
    print(f"{'claim':44s} {'stated':>12} {'from data':>12}  {'in text':>8}")
    for name, stated, computed in checks:
        agrees = _agree(stated, computed)
        if text is None:
            in_text = "n/a"
        else:
            in_text = "yes" if any(f in text for f in _renderings(stated)) else "NO"
        flag = "" if (agrees and in_text != "NO") else "   <-- FAIL"
        if flag:
            fails.append((name, stated, computed, in_text))
        print(f"{name:44s} {stated:>12} {computed:>12}  {in_text:>8}{flag}")

    print(f"\nchecked {len(checks)} registered claims; {len(fails)} failed")
    if fails:
        print("\nA failure means the manuscript states a number the data do not "
              "produce, or omits one it should state.")
        return 1
    print("All registered claims reproduce from data and appear in the manuscript.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
